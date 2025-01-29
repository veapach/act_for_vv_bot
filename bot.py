import os
from io import BytesIO
from aiogram import Bot, Dispatcher, Router
from aiogram.types import ContentType, Message, ReplyKeyboardMarkup, KeyboardButton
from aiogram.fsm.storage.memory import MemoryStorage
from aiogram.types import FSInputFile
from aiogram import F
from aiogram.types import Message
from aiogram.client.session.aiohttp import AiohttpSession
from docx import Document
from docx.shared import Cm
from docx.oxml import parse_xml
from docx.oxml.ns import nsmap
from PIL import Image
from config import API_TOKEN
import asyncio


# –°–æ–∑–¥–∞–µ–º –∫–ª–∞–≤–∏–∞—Ç—É—Ä—É —Å –∫–Ω–æ–ø–∫–æ–π "–ì–æ—Ç–æ–≤–æ ‚úÖ"
done_button = KeyboardButton(text="–ì–æ—Ç–æ–≤–æ ‚úÖ")
keyboard = ReplyKeyboardMarkup(keyboard=[[done_button]], resize_keyboard=True)

# –°–æ–∑–¥–∞–µ–º –±–æ—Ç–∞ –∏ –¥–∏—Å–ø–µ—Ç—á–µ—Ä
bot = Bot(token=API_TOKEN, session=AiohttpSession())
dp = Dispatcher(storage=MemoryStorage())
router = Router()

# –•—Ä–∞–Ω–∏–ª–∏—â–µ –¥–ª—è —Ñ–æ—Ç–æ–≥—Ä–∞—Ñ–∏–π –∏ –¥–∞–Ω–Ω—ã—Ö –ø–æ–ª—å–∑–æ–≤–∞—Ç–µ–ª–µ–π
user_photos = {}
user_data = {}


@router.message(F.text == "/start")
async def start_handler(message: Message):
    await message.reply(
        "–ü—Ä–∏–≤–µ—Ç! –û—Ç–ø—Ä–∞–≤—å –º–Ω–µ –Ω–µ—Å–∫–æ–ª—å–∫–æ —Ñ–æ—Ç–æ, –∏ —è –≤—Å—Ç–∞–≤–ª—é –∏—Ö –≤ –æ—Ç—á–µ—Ç. –ö–æ–≥–¥–∞ –∑–∞–∫–æ–Ω—á–∏—à—å, –Ω–∞–∂–º–∏ –∫–Ω–æ–ø–∫—É <b>–ì–æ—Ç–æ–≤–æ</b>.",
        parse_mode="HTML",
    )
    username = message.from_user.username or "–Ω–µ–∏–∑–≤–µ—Å—Ç–Ω—ã–π"
    print(f"[LOG] - –ó–∞–ø—Ä–æ—Å –æ—Ç –Ω–æ–≤–æ–≥–æ –ø–æ–ª—å–∑–æ–≤–∞—Ç–µ–ª—è! @{username}")


@router.message(F.photo)
async def photo_handler(message: Message):
    user_id = message.from_user.id

    if user_id not in user_photos:
        user_photos[user_id] = []

    # –°–æ—Ö—Ä–∞–Ω—è–µ–º —Ç–æ–ª—å–∫–æ —Ñ–æ—Ç–æ —Å –Ω–∞–∏–≤—ã—Å—à–∏–º —Ä–∞–∑—Ä–µ—à–µ–Ω–∏–µ–º
    highest_quality_photo = max(message.photo, key=lambda p: p.file_size)
    user_photos[user_id].append(highest_quality_photo.file_id)

    # –û—Ç–ø—Ä–∞–≤–ª—è–µ–º —Å–æ–æ–±—â–µ–Ω–∏–µ —Å –∫–Ω–æ–ø–∫–æ–π "–ì–æ—Ç–æ–≤–æ ‚úÖ" —Ç–æ–ª—å–∫–æ –æ–¥–∏–Ω —Ä–∞–∑
    if len(user_photos[user_id]) == 1:
        await message.reply(
            "üì∏ –§–æ—Ç–æ –ø–æ–ª—É—á–µ–Ω–æ! –û—Ç–ø—Ä–∞–≤—å –µ—â—ë –∏–ª–∏ –Ω–∞–∂–º–∏ –∫–Ω–æ–ø–∫—É <b>–ì–æ—Ç–æ–≤–æ</b>, —á—Ç–æ–±—ã –∑–∞–≤–µ—Ä—à–∏—Ç—å.",
            reply_markup=keyboard,
            parse_mode="HTML",
        )
        username = message.from_user.username or "–Ω–µ–∏–∑–≤–µ—Å—Ç–Ω—ã–π"
        print(f"[LOG] - –ü–æ–ª—å–∑–æ–≤–∞—Ç–µ–ª—å –æ—Ç–ø—Ä–∞–≤–∏–ª —Ñ–æ—Ç–æ @{username}")


@router.message(F.text == "–ì–æ—Ç–æ–≤–æ ‚úÖ")
async def done_handler(message: Message):
    user_id = message.from_user.id

    if user_id not in user_photos or not user_photos[user_id]:
        await message.reply(
            "‚ö†Ô∏è –í—ã –µ—â—ë –Ω–µ –æ—Ç–ø—Ä–∞–≤–∏–ª–∏ –Ω–∏ –æ–¥–Ω–æ–π —Ñ–æ—Ç–æ–≥—Ä–∞—Ñ–∏–∏.", parse_mode="HTML"
        )
        return

    await message.reply(
        "üìÖ –ü–æ–∂–∞–ª—É–π—Å—Ç–∞, –≤–≤–µ–¥–∏—Ç–µ –¥–∞—Ç—É –≤ —Ñ–æ—Ä–º–∞—Ç–µ –î–î.–ú–ú.–ì–ì–ì–ì:", parse_mode="HTML"
    )
    username = message.from_user.username or "–Ω–µ–∏–∑–≤–µ—Å—Ç–Ω—ã–π"
    print(f"[LOG] - –§–æ—Ç–æ –æ—Ç –ø–æ–ª—å–∑–æ–≤–∞—Ç–µ–ª—è —Å–æ—Ö—Ä–∞–Ω–µ–Ω—ã @{username}")

    # –ò–Ω–∏—Ü–∏–∞–ª–∏–∑–∏—Ä—É–µ–º –¥–∞–Ω–Ω—ã–µ –ø–æ–ª—å–∑–æ–≤–∞—Ç–µ–ª—è
    user_data[user_id] = {"photos": user_photos[user_id]}


@router.message(F.text.regexp(r"\d{2}\.\d{2}\.\d{4}"))
async def date_handler(message: Message):
    user_id = message.from_user.id

    if user_id not in user_data:
        await message.reply(
            "‚ö†Ô∏è –ü–æ–∂–∞–ª—É–π—Å—Ç–∞, —Å–Ω–∞—á–∞–ª–∞ –Ω–∞–∂–º–∏—Ç–µ '–ì–æ—Ç–æ–≤–æ ‚úÖ'.", parse_mode="HTML"
        )
        return

    user_data[user_id]["date"] = message.text
    await message.reply("üìç –ü–æ–∂–∞–ª—É–π—Å—Ç–∞, –≤–≤–µ–¥–∏—Ç–µ –∞–¥—Ä–µ—Å:", parse_mode="HTML")
    username = message.from_user.username or "–Ω–µ–∏–∑–≤–µ—Å—Ç–Ω—ã–π"
    print(f"[LOG] - –î–∞—Ç–∞ –æ—Ç –ø–æ–ª—å–∑–æ–≤–∞—Ç–µ–ª—è —Å–æ—Ö—Ä–∞–Ω–µ–Ω–∞ @{username}")


@router.message(F.text)
async def address_handler(message: Message):
    user_id = message.from_user.id

    if user_id not in user_data or "date" not in user_data[user_id]:
        return

    user_data[user_id]["address"] = message.text
    await message.reply("üìù –°–æ–∑–¥–∞—é –¥–æ–∫—É–º–µ–Ω—Ç, –ø–æ–¥–æ–∂–¥–∏—Ç–µ –Ω–µ–º–Ω–æ–≥–æ...", parse_mode="HTML")
    username = message.from_user.username or "–Ω–µ–∏–∑–≤–µ—Å—Ç–Ω—ã–π"
    print(
        f"[LOG] - –ê–¥—Ä–µ—Å –æ—Ç –ø–æ–ª—å–∑–æ–≤–∞—Ç–µ–ª—è —Å–æ—Ö—Ä–∞–Ω–µ–Ω, –Ω–∞—á–∞–ª–∞—Å—å –æ–±—Ä–∞–±–æ—Ç–∫–∞ –¥–æ–∫—É–º–µ–Ω—Ç–∞ @{username}"
    )

    try:
        # –ì–µ–Ω–µ—Ä–∞—Ü–∏—è –¥–æ–∫—É–º–µ–Ω—Ç–∞ Word
        output_file = await generate_document(user_id, user_data[user_id])

        # –û—Ç–ø—Ä–∞–≤–∫–∞ –¥–æ–∫—É–º–µ–Ω—Ç–∞ –ø–æ–ª—å–∑–æ–≤–∞—Ç–µ–ª—é
        await message.answer_document(FSInputFile(output_file))
        print(f"[LOG] - –ì–æ—Ç–æ–≤—ã–π –æ—Ç—á–µ—Ç –æ—Ç–ø—Ä–∞–≤–ª–µ–Ω –ø–æ–ª—å–∑–æ–≤–∞—Ç–µ–ª—é @{username}")

        # –û—á–∏—Å—Ç–∫–∞ –≤—Ä–µ–º–µ–Ω–Ω—ã—Ö –¥–∞–Ω–Ω—ã—Ö
        os.remove(output_file)
        del user_photos[user_id]
        del user_data[user_id]
    except Exception as e:
        await message.reply(f"‚ùå –ü—Ä–æ–∏–∑–æ—à–ª–∞ –æ—à–∏–±–∫–∞: {e}", parse_mode="HTML")


# –ì–µ–Ω–µ—Ä–∞—Ü–∏—è –¥–æ–∫—É–º–µ–Ω—Ç–∞ Word
async def generate_document(user_id, user_info):
    doc = Document("template.docx")

    # –í—Å—Ç–∞–≤–∫–∞ –¥–∞—Ç—ã –∏ –∞–¥—Ä–µ—Å–∞ –≤ —Ç–∞–±–ª–∏—Ü—ã –¥–æ–∫—É–º–µ–Ω—Ç–∞
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if "[–¥–∞—Ç–∞]" in cell.text:
                    cell.text = cell.text.replace("[–¥–∞—Ç–∞]", user_info["date"])
                if "[–∞–¥—Ä–µ—Å]" in cell.text:
                    cell.text = cell.text.replace("[–∞–¥—Ä–µ—Å]", user_info["address"])
                if "[–≤—Å—Ç–∞–≤–∫–∞]" in cell.text:
                    cell.text = ""  # –£–¥–∞–ª—è–µ–º —Ç–µ–∫—Å—Ç [–≤—Å—Ç–∞–≤–∫–∞]
                    for photo_id in user_info["photos"]:
                        photo = await bot.download(file=await bot.get_file(photo_id))
                        image = Image.open(BytesIO(photo.read()))

                        # –ò–∑–º–µ–Ω–µ–Ω–∏–µ —Ä–∞–∑–º–µ—Ä–∞
                        width_cm = 18
                        height_cm = 13.5
                        dpi = 96
                        width_px = int((width_cm / 2.54) * dpi)
                        height_px = int((height_cm / 2.54) * dpi)
                        image.thumbnail((width_px, height_px))

                        # –°–æ—Ö—Ä–∞–Ω–µ–Ω–∏–µ –≤—Ä–µ–º–µ–Ω–Ω–æ–≥–æ –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏—è
                        temp_image_path = f"temp_{user_id}_{photo_id}.jpg"
                        image.save(temp_image_path)

                        # –í—Å—Ç–∞–≤–∫–∞ –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏—è
                        paragraph = cell.add_paragraph()
                        paragraph.alignment = 1  # –¶–µ–Ω—Ç—Ä–∏—Ä–æ–≤–∞–Ω–∏–µ –ø–∞—Ä–∞–≥—Ä–∞—Ñ–∞
                        run = paragraph.add_run()
                        run.add_picture(
                            temp_image_path, width=Cm(width_cm), height=Cm(height_cm)
                        )

                        os.remove(temp_image_path)

    # –§–æ—Ä–º–∏—Ä–æ–≤–∞–Ω–∏–µ –∏–º–µ–Ω–∏ –≤—ã—Ö–æ–¥–Ω–æ–≥–æ —Ñ–∞–π–ª–∞
    date = user_info["date"]
    address = user_info["address"]
    output_path = f"–û—Ç—á–µ—Ç –æ –ø—Ä–æ–≤–µ–¥–µ–Ω–Ω–æ–º –¢–û-{date}  {address}.docx"
    doc.save(output_path)
    return output_path


async def main():
    dp.include_router(router)
    await bot.delete_webhook(drop_pending_updates=True)
    print("[LOG] - –ë–æ—Ç –∑–∞–ø—É—â–µ–Ω —É—Å–ø–µ—à–Ω–æ")
    await dp.start_polling(bot, parse_mode="HTML")


if __name__ == "__main__":
    asyncio.run(main())
