from docx import Document
from docx.shared import Cm
from PIL import Image
from io import BytesIO
import os
from bot_instance import bot
from database import Database

db = Database()


async def generate_document(user_id, user_info):
    doc = Document("template.docx")

    user_record = await db.get_user(user_id)
    if user_record:
        first_name, last_name = user_record
        full_name = f"{last_name} {first_name}"
    else:
        full_name = "ФИО не указано"

    user_info = user_info.copy()

    if isinstance(user_info.get("works", ""), list):
        user_info["works"] = "\n• " + "\n• ".join(user_info["works"])

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if "[дата]" in cell.text:
                    cell.text = cell.text.replace("[дата]", user_info["date"])
                if "[адрес]" in cell.text:
                    cell.text = cell.text.replace("[адрес]", user_info["address"])
                if "[классификация]" in cell.text:
                    cell.text = cell.text.replace(
                        "[классификация]", user_info["classification"]
                    )
                if "[вставка]" in cell.text:
                    cell.text = ""
                    for photo_id in user_info["photos"]:
                        photo = await bot.download(file=await bot.get_file(photo_id))
                        image = Image.open(BytesIO(photo.read()))

                        width_cm = 18
                        height_cm = 13.5
                        dpi = 96
                        width_px = int((width_cm / 2.54) * dpi)
                        height_px = int((height_cm / 2.54) * dpi)
                        image.thumbnail((width_px, height_px))

                        temp_image_path = f"temp_{user_id}_{photo_id}.jpg"
                        image.save(temp_image_path)

                        paragraph = cell.add_paragraph()
                        paragraph.alignment = 1
                        run = paragraph.add_run()
                        run.add_picture(
                            temp_image_path, width=Cm(width_cm), height=Cm(height_cm)
                        )

                        os.remove(temp_image_path)

                if "Печь [классификация]" in cell.text:
                    classification = user_info.get("classification", "")
                    if classification:
                        cell.text = cell.text.replace(
                            "Печь [классификация]", classification
                        )
                    else:
                        cell.text = cell.text.replace("[классификация]", "")

                if "[работы]" in cell.text:
                    works = user_info.get("works", "Работы не проводились")
                    cell.text = cell.text.replace("[работы]", works)

                if "[материалы]" in cell.text:
                    materials = user_info.get("materials", "")
                    cell.text = cell.text.replace("[материалы]", materials)

                if "[рекомендации]" in cell.text:
                    recommendations = user_info.get("recommendations", "")
                    cell.text = cell.text.replace("[рекомендации]", recommendations)

                if "[дефекты]" in cell.text:
                    defects = user_info.get("defects", "")
                    cell.text = cell.text.replace("[дефекты]", defects)

                if "[доп_работы]" in cell.text:
                    additional_works = user_info.get("additional_works", "")
                    cell.text = cell.text.replace("[доп_работы]", additional_works)

                if "[фио]" in cell.text:
                    cell.text = cell.text.replace("[фио]", full_name)

                if "[комментарии]" in cell.text:
                    comments = user_info.get("comments", "")
                    cell.text = cell.text.replace("[комментарии]", comments)

    date = user_info["date"]
    address = user_info["address"]
    output_path = f"Акт выполненных работ {date}  {address}.docx"
    doc.save(output_path)
    return output_path
